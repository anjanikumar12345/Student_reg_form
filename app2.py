import streamlit as st
import pandas as pd
from datetime import datetime
import requests
from googletrans import Translator # type: ignore
from docx import Document





st.markdown("""
<div style="background-color:#ffcc00; padding:20px; border-radius:10px;">
  <marquee behavior="scroll" direction="left" scrollamount="6">
    📢 DIVINE GRACE ASSEMBLY ROAS NO.3 , Vishwanadha Colony , Shambunipet , Warangal , 506005
  </marquee>
</div>
""", unsafe_allow_html=True)



# Set up translator
translator = Translator()

# Split KJV Bible books by testament
old_testament = [
    'Genesis', 'Exodus', 'Leviticus', 'Numbers', 'Deuteronomy', 'Joshua', 'Judges',
    'Ruth', '1 Samuel', '2 Samuel', '1 Kings', '2 Kings', '1 Chronicles', '2 Chronicles',
    'Ezra', 'Nehemiah', 'Esther', 'Job', 'Psalms', 'Proverbs', 'Ecclesiastes', 'Song of Solomon',
    'Isaiah', 'Jeremiah', 'Lamentations', 'Ezekiel', 'Daniel', 'Hosea', 'Joel', 'Amos',
    'Obadiah', 'Jonah', 'Micah', 'Nahum', 'Habakkuk', 'Zephaniah', 'Haggai', 'Zechariah', 'Malachi'
]

new_testament = [
    'Matthew', 'Mark', 'Luke', 'John', 'Acts', 'Romans', '1 Corinthians', '2 Corinthians',
    'Galatians', 'Ephesians', 'Philippians', 'Colossians', '1 Thessalonians', '2 Thessalonians',
    '1 Timothy', '2 Timothy', 'Titus', 'Philemon', 'Hebrews', 'James', '1 Peter', '2 Peter',
    '1 John', '2 John', '3 John', 'Jude', 'Revelation'
]

# Title and Image
st.set_page_config(page_title="Divine Grace Assembly", layout="wide")
st.image("C:/Users/vanja/Downloads/IMG_20250630_120426.jpg", width=350)

st.markdown("""
    <style>
        body {
            background-image: url('"C:/Users/vanja/Downloads/istockphoto-2170275560-1024x1024.jpg"');
            background-size: cover;
            background-position: center;
            background-repeat: no-repeat;
        }
        .main {
            background-color: rgba(255, 255, 255, 0.9);
            padding: 2rem;
            border-radius: 10px;
            box-shadow: 0px 4px 6px rgba(0,0,0,0.1);
        }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: center; color: #2c3e50;'>Divine Grace Assembly</h1>", unsafe_allow_html=True)
st.markdown("<h2 style='text-align: center; color: #8e44ad;'>Pastor Ramesh Paul - Founder & Director</h2>", unsafe_allow_html=True)

st.title("📖 Sunday Bible Diary")

# --- Bible Lookup Section ---
st.header("🔍 King James Bible (KJV) Verse Lookup")
testament = st.radio("టెస్టమెంట్ ఎంచుకోండి (Select Testament)", ["పాత ఒడంబడిక (Old Testament)", "క్రొత్త ఒడంబడిక (New Testament)"])
if testament.startswith("పాత") or testament == "Old Testament":
    book = st.selectbox("పుస్తకం ఎంచుకోండి (Select Book)", old_testament)
else:
    book = st.selectbox("పుస్తకం ఎంచుకోండి (Select Book)", new_testament)

chapter = st.number_input("Chapter", min_value=1, step=1)
verse = st.number_input("Verse", min_value=1, step=1)

if st.button("Get Verse"):
    if book and chapter and verse:
        url = f"https://bible-api.com/{book}%20{chapter}:{verse}?translation=kjv"
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            eng_verse = data.get("text", "Verse not found").strip()
            st.success(f"{book} {chapter}:{verse} - {eng_verse}")
            telugu_verse = translator.translate(eng_verse, src='en', dest='te').text
            st.info(f"Telugu Translation: {telugu_verse}")
        else:
            st.error("Verse not found. Please check inputs.")

# --- Diary entry form ---
with st.form("entry_form"):
    st.subheader("Add Sunday Verse")
    date = st.date_input("Date")
    time = st.time_input("Time")
    verse_english = st.text_area("Bible Verse (English)")
    verse_telugu = st.text_area("బైబిల్ వాక్యం (Telugu)")
    submit = st.form_submit_button("✅ Save Entry")

# Initialize storage if not available
if "diary" not in st.session_state:
    st.session_state.diary = []

# Save data
if submit:
    st.session_state.diary.append({
        "Date": date.strftime("%Y-%m-%d"),
        "Time": time.strftime("%H:%M"),
        "Verse (English)": verse_english,
        "Verse (Telugu)": verse_telugu
    })
    st.success("Entry saved successfully!")

# Display stored data
if st.session_state.diary:
    st.subheader("📜 Saved Bible Verses")
    df = pd.DataFrame(st.session_state.diary)
    st.dataframe(df)

    # Delete entry
    delete_index = st.number_input("Enter row index to delete (starting from 0)", min_value=0, max_value=len(df)-1, step=1)
    if st.button("🗑️ Delete Entry"):
        st.session_state.diary.pop(delete_index)
        st.success("Entry deleted successfully!")

    # Download CSV
    csv = pd.DataFrame(st.session_state.diary).to_csv(index=False).encode('utf-8')
    st.download_button("📥 Download Diary as CSV", csv, "divine_grace_diary.csv", "text/csv")

    # Download DOCX
    if st.button("📄 Save Entries as Word Document"):
        doc = Document()
        doc.add_heading("Divine Grace Assembly - Bible Diary", 0)
        for entry in st.session_state.diary:
            doc.add_paragraph(f"Date: {entry['Date']}")
            doc.add_paragraph(f"Time: {entry['Time']}")
            doc.add_paragraph(f"Verse (English): {entry['Verse (English)']}")
            doc.add_paragraph(f"Verse (Telugu): {entry['Verse (Telugu)']}")
            doc.add_paragraph("-" * 50)
        doc_path =   doc_path = "C:\\Users\\vanja\\OneDrive\\Desktop\\reg_form\\Divine_Grace_Bible_Diary.docx"

        doc.save(doc_path)

        doc.save(doc_path)
        with open(doc_path, "rb") as f:
            st.download_button("📥 Download as Word Document", f, file_name="Divine_Grace_Bible_Diary.docx")


st.markdown("""
    <style>
    .stApp {
        background-color: #bce5f3; /* Light blue background */
    }
    .stButton > button {
        background-color: #f4fc0b ;
        color: white;
        font-size: 20px;
        border-radius: 10px;
    }
    </style>
""", unsafe_allow_html=True)







st.markdown("""
<div style="background-color:#ffcc00; padding:20px; border-radius:10px;">
  <marquee behavior="scroll" direction="left" scrollamount="6">
    📢 DIVINE GRACE ASSEMBLY ROAS NO.3 , Vishwanadha Colony , Shambunipet , Warangal , 506005
  </marquee>
</div>
""", unsafe_allow_html=True)